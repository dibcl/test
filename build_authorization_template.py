from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = r"C:\Users\Administrator\Desktop\zte-research\云桌面Debian兼容性与Guest管理安全测试授权书范本.docx"
INK = "1F2937"
BLUE = "1F4E78"
LIGHT = "EAF1F7"
GRAY = "F2F4F7"
MUTED = "667085"
RED = "9B1C1C"
GOLD = "FFF4CE"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        unique_tcs = []
        for cell in row.cells:
            if all(cell._tc is not existing for existing in unique_tcs):
                unique_tcs.append(cell._tc)
        if len(unique_tcs) == 1:
            tc_pr = unique_tcs[0].get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(sum(widths_dxa)))
            tc_w.set(qn("w:type"), "dxa")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def set_font(run, name="宋体", size=10.5, bold=False, color=INK, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_para(p, before=0, after=6, line=1.10, align=None, keep=False):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if keep:
        p.paragraph_format.keep_with_next = True


def add_para(doc, text="", bold_prefix=None, after=6, color=INK, size=10.5, italic=False):
    p = doc.add_paragraph()
    set_para(p, after=after)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True, size=size, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, size=size, color=color, italic=italic)
    else:
        r = p.add_run(text)
        set_font(r, size=size, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_kv_table(doc, rows, widths=(2100, 7260), header=None):
    if header:
        hp = doc.add_paragraph()
        set_para(hp, before=3, after=4, keep=True)
        set_font(hp.add_run(header), name="微软雅黑", size=11, bold=True, color=BLUE)
    total_rows = len(rows)
    table = doc.add_table(rows=total_rows, cols=2)
    table.style = "Table Grid"
    idx = 0
    for label, value in rows:
        left, right = table.rows[idx].cells
        set_cell_shading(left, GRAY)
        lp = left.paragraphs[0]
        rp = right.paragraphs[0]
        set_para(lp, after=0)
        set_para(rp, after=0)
        set_font(lp.add_run(label), bold=True)
        set_font(rp.add_run(value))
        idx += 1
    set_table_geometry(table, list(widths))
    mark_header_row(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_scope_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(("授权选项", "测试项目", "授权范围/限制")):
        set_cell_shading(hdr[i], LIGHT)
        p = hdr[i].paragraphs[0]
        set_para(p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_font(p.add_run(text), bold=True, color=BLUE)
    for checked, title, detail in rows:
        cells = table.add_row().cells
        vals = (checked, title, detail)
        for i, val in enumerate(vals):
            p = cells[i].paragraphs[0]
            set_para(p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT)
            set_font(p.add_run(val), bold=(i == 1))
    set_table_geometry(table, [1080, 2520, 5760])
    mark_header_row(table.rows[0])
    return table


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for e in (fld_begin, instr, fld_sep, text, fld_end):
        run._r.append(e)
    set_font(run, size=9, color=MUTED)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.82)
sec.bottom_margin = Inches(0.78)
sec.left_margin = Inches(1.0)
sec.right_margin = Inches(1.0)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "宋体"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for idx, size, before, after in ((1, 15, 14, 7), (2, 12.5, 10, 5), (3, 11.5, 8, 4)):
    st = styles[f"Heading {idx}"]
    st.font.name = "微软雅黑"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(BLUE if idx < 3 else INK)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

header = sec.header.paragraphs[0]
set_para(header, after=0)
set_font(header.add_run("受控安全测试文件 | 未签署前无授权效力"), size=8.5, color=MUTED)
footer = sec.footer.paragraphs[0]
set_para(footer, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
set_font(footer.add_run("云桌面 Debian 兼容性与 Guest 管理安全测试授权书范本  |  第 "), size=8.5, color=MUTED)
add_page_field(footer)
set_font(footer.add_run(" 页"), size=8.5, color=MUTED)

p = doc.add_paragraph()
set_para(p, before=8, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(p.add_run("云桌面 Debian 兼容性与 Guest 管理"), name="微软雅黑", size=21, bold=True, color="000000")
p = doc.add_paragraph()
set_para(p, after=14, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(p.add_run("安全测试授权书"), name="微软雅黑", size=21, bold=True, color="000000")
p = doc.add_paragraph()
set_para(p, after=16, align=WD_ALIGN_PARAGRAPH.CENTER)
set_font(p.add_run("（含受控 Guest 身份、管理协议与遥测仿真测试）"), name="微软雅黑", size=11.5, color=MUTED)

add_kv_table(doc, [
    ("文件编号", "[填写：项目编号/工单号/变更单号]"),
    ("版本", "V1.0（签署时填写最终版本）"),
    ("授权单位", "[填写完整法人名称及统一社会信用代码]"),
    ("被授权单位/人员", "[填写公司、部门、姓名、工号及联系方式]"),
    ("有效期", "自 [年/月/日 时:分] 至 [年/月/日 时:分]（含时区）"),
    ("保密等级", "[公开/内部/秘密/其他：________]"),
], header="文件控制信息")

p = doc.add_paragraph()
set_para(p, before=4, after=8)
set_cell_text = p.add_run("重要提示：")
set_font(set_cell_text, bold=True, color=RED)
set_font(p.add_run("本文件仅为授权书范本，不构成法律意见；必须由资产所有人、平台运营方、网络安全/法务及测试负责人补全所有字段并正式签署。未填写实例标识、测试窗口、后台联系人和恢复责任，或仅有个人签名时，不得据此开展身份/遥测仿真、协议重放、DD、分区或启动链修改。"), color=RED)

add_heading(doc, "一、授权背景与目的", 1)
add_para(doc, "为验证指定云桌面平台在 Guest 操作系统由 Windows 变更为 Debian/Linux 后的硬件兼容性、生命周期管理、状态识别、客户端连接、Agent 检测、心跳及资产遥测能力，并评估现有检测机制在受控仿真条件下的准确性和抗规避能力，授权方同意在本文件明确列出的隔离测试资产和时间窗口内开展安全测试。")
add_para(doc, "测试成果仅用于兼容性开发、安全评估、缺陷修复和内部验收，不得用于公众环境长期规避平台策略、计费、访问控制或安全监测。")

add_heading(doc, "二、授权主体与责任人", 1)
add_kv_table(doc, [
    ("资产所有/运营单位", "[完整法人名称]"),
    ("业务主管", "[姓名/职务/电话/企业邮箱]"),
    ("平台技术负责人", "[姓名/职务/电话/企业邮箱]"),
    ("网络安全负责人", "[姓名/职务/电话/企业邮箱]"),
    ("法务/合规负责人", "[姓名/职务/电话/企业邮箱]"),
    ("后台观察人", "[姓名/值守时间/电话/企业邮箱]"),
    ("测试负责人", "[姓名/单位/工号/电话/企业邮箱]"),
    ("应急恢复负责人", "[姓名/7×24 联系方式/恢复权限说明]"),
])

add_heading(doc, "三、指定测试资产（必须精确到实例）", 1)
add_kv_table(doc, [
    ("产品/SKU", "[中国移动云电脑公众版具体套餐/版本]"),
    ("账号/订单号", "[填写；必要时部分脱敏，但签署原件须可唯一定位]"),
    ("实例 UUID/VMID", "[填写完整标识]"),
    ("区域/资源池/宿主类型", "[填写]"),
    ("公网/管理 IP", "[填写授权测试地址范围]"),
    ("原始镜像", "[Windows 镜像名称、版本、镜像 ID]"),
    ("目标系统", "[Debian 版本、架构、镜像 SHA-256]"),
    ("系统盘/数据盘", "[磁盘号、容量、序列/卷标及允许改写范围]"),
    ("关联后台系统", "[仅列本次允许观察/验证的系统名称]"),
    ("禁止触及资产", "除上述实例外的全部账号、实例、Host、管理节点及其他用户数据"),
])
add_para(doc, "若本节任一关键标识为空，涉及主动协议、身份或遥测仿真的授权自动无效。", color=RED, italic=True)

add_heading(doc, "四、测试窗口与环境控制", 1)
add_kv_table(doc, [
    ("授权开始", "[年/月/日 时:分，时区]"),
    ("授权结束", "[年/月/日 时:分，时区]"),
    ("维护窗口", "[填写具体窗口]"),
    ("环境性质", "□ 隔离测试资源池  □ 可牺牲测试实例  □ 生产同构但无生产业务  □ 其他：_____"),
    ("后台监看", "□ 全程在线  □ 关键阶段在线  □ 不允许主动仿真（若无后台监看）"),
    ("最大测试速率", "[消息/秒、重启次数、并发数及流量上限]"),
    ("异常告警抑制", "[是否批准、范围及负责人员；不得默认关闭]"),
])

add_heading(doc, "五、授权操作范围（逐项勾选，未勾选即不授权）", 1)
add_scope_table(doc, [
    ("□ 允许", "只读基线调查", "文件、注册表、服务、驱动、进程、日志、网络连接元数据、PE/签名/哈希和 VirtIO 映射分析。"),
    ("□ 允许", "系统与磁盘变更", "DD、分区调整、GRUB/Legacy Boot/BCD 修改、Debian 安装、重启/关机；允许改写范围须与第三节一致。"),
    ("□ 允许", "Linux 兼容组件开发", "编写、安装和调试仅用于本实例的 Guest transport、QGA、状态、power/reboot 和 inventory 兼容组件。"),
    ("□ 允许", "Guest 身份仿真", "在测试窗口内模拟 Windows Guest/Agent 身份，用于验证后台识别与抗规避能力。必须由后台观察人同步记录。"),
    ("□ 允许", "固定 OS/资产遥测", "返回经审批的合成 Windows OS、软件清单、KB、主机名、版本等测试数据；测试数据集见附件。"),
    ("□ 允许", "Heartbeat/状态仿真", "构造或复现实例范围内的 Agent 在线、健康、心跳和状态上报，仅可使用批准频率与测试身份。"),
    ("□ 允许", "私有协议主动测试", "对指定 VirtIO Serial/localhost 管理通道进行消息构造、回放、模糊测试或兼容实现；允许消息类型和速率见附件。"),
    ("□ 允许", "实例身份材料使用", "仅在原实例内临时使用该实例已经分配的密钥/证书/UUID；禁止导出到其他实例或伪造其他资产。"),
    ("□ 允许", "官方客户端验证", "验证 ICE/RAP/Vdagent、显示、输入、音频、USB/文件/打印重定向及会话恢复。"),
    ("□ 允许", "恢复与回退", "恢复官方 Windows 镜像、系统盘快照、Agent 组件和平台状态；由应急恢复负责人执行或监护。"),
])

add_heading(doc, "六、受控身份与遥测仿真专项授权", 1)
add_para(doc, "仅当第五条对应项目已勾选且本条全部填写时，测试人员方可开展相关测试。")
add_kv_table(doc, [
    ("仿真目的", "[例：验证后台能否识别 OS 替换、静态 inventory 和 Agent 行为异常]"),
    ("允许仿真的身份", "[仅本实例 Windows Guest/指定 Agent 名称与版本]"),
    ("批准的合成数据集", "[附件编号、字段列表、固定值/变化规则及数据负责人]"),
    ("允许消息类别", "[心跳/状态/OS inventory/软件 inventory/power response/其他]"),
    ("禁止消息类别", "[身份注册、计费、授权、跨租户、文件下发、远程命令等]"),
    ("允许的重放范围", "[无/仅本实例同一窗口内/指定消息及次数]"),
    ("凭据处理", "[可用材料、存储位置、访问人、销毁时间；禁止跨实例复制]"),
    ("后台预期结果", "[应识别/应告警/允许短暂误判/其他验收标准]"),
    ("强制销毁时间", "授权结束后 [__] 小时内删除仿真组件、合成数据和临时凭据"),
])

add_heading(doc, "七、明确禁止事项", 1)
for text in [
    "7.1 不得访问、识别、复制或影响其他用户、租户、实例、Host、管理节点或生产业务数据。",
    "7.2 不得扩大扫描或消息发送范围，不得绕过账号、计费、许可、访问控制或租户隔离。",
    "7.3 不得实施拒绝服务、资源耗尽、跨实例凭据复用、持久化后门或未经批准的远程命令执行。",
    "7.4 不得在授权窗口之外继续运行身份/遥测仿真组件，不得将仿真组件用于长期规避平台管理。",
    "7.5 不得向第三方公开可复现的敏感协议细节、密钥、实例身份或未修复漏洞；披露流程见第十一条。",
    "7.6 未经应急恢复负责人确认，不得在缺少可用快照、官方重装或人工恢复保障时改写唯一系统盘。",
]:
    add_para(doc, text, after=4)

add_heading(doc, "八、停止条件与应急响应", 1)
add_kv_table(doc, [
    ("立即停止条件", "出现跨实例影响、异常流量、后台不可控告警、实例锁定、数据泄露、Host 异常、恢复路径失效或负责人要求停止。"),
    ("停止动作", "停止主动消息；保存必要证据；通知后台观察人与安全负责人；不得擅自扩大测试或尝试绕过锁定。"),
    ("恢复顺序", "[平台解除测试状态] → [恢复系统盘/镜像] → [恢复官方 Agent] → [验证管理与客户端] → [关闭工单]。"),
    ("恢复时限/SLA", "授权方承诺在触发停止条件后 [__] 分钟内响应、[__] 小时内完成恢复或提供替代实例。"),
    ("不可恢复责任", "因严格按已批准方案执行导致的实例重装、锁定或数据损失，由 [授权单位/部门] 承担恢复与业务协调责任。"),
])

add_heading(doc, "九、测试计划、成功标准与变更控制", 1)
add_para(doc, "9.1 测试必须按附件《测试计划》分阶段执行：基线采集、离线镜像准备、可回滚启动、Debian 硬件验证、Guest 管理验证、受控仿真、稳定性、恢复。")
add_para(doc, "9.2 每阶段须定义输入、操作、期望结果、后台观察项、停止条件和回滚动作；前一阶段未通过不得自动进入下一阶段。")
add_para(doc, "9.3 成功标准至少包括：Debian 正常启动；VirtIO 磁盘/网卡/串口正常；SSH 正常；平台 power state 正确；授权范围内的 Guest 状态测试达到预期；完成恢复并清除临时组件。")
add_para(doc, "9.4 超出本授权书或附件的消息类型、凭据使用、实例、时间窗口或破坏性操作，必须重新提交书面变更并由原审批角色再次签署。")

add_heading(doc, "十、证据、数据与保密", 1)
add_kv_table(doc, [
    ("允许采集", "本实例测试日志、时间线、进程/服务/驱动状态、网络连接元数据、消息类别与长度、后台状态截图、镜像/组件哈希。"),
    ("默认脱敏", "账号、手机号、UUID、VMID、IP、密钥、token、证书私钥、其他用户信息及业务数据。"),
    ("证据存储", "[批准位置、访问控制、加密方式]"),
    ("保留期限", "测试结束后 [__] 天；到期由 [负责人] 审核销毁。"),
    ("报告接收人", "[开发/安全/平台/法务名单]"),
])

add_heading(doc, "十一、缺陷报告与披露", 1)
add_para(doc, "发现平台识别、Agent 检测、身份校验、心跳、inventory 或租户隔离缺陷时，应在 [__] 小时内通过 [工单/邮箱/漏洞平台] 报告。未经书面同意，不得公开漏洞细节、可运行仿真组件或复现步骤。授权方应在 [__] 个工作日内确认接收并给出修复/缓解计划。")

add_heading(doc, "十二、授权效力与解释", 1)
add_para(doc, "12.1 本授权仅对第三节指定资产、第四节时间窗口和第五/六条勾选项目有效。")
add_para(doc, "12.2 本授权不当然豁免适用法律、监管要求、用户协议、数据保护义务和授权单位内部制度；冲突时以更严格要求为准。")
add_para(doc, "12.3 电子签名、盖章扫描件或经授权单位认可的工单审批记录可作为附件，与本授权书共同构成完整授权证据。")
add_para(doc, "12.4 未经全部必需审批人签署，本文件仅为草案，不产生主动安全测试授权效力。", color=RED)

add_heading(doc, "十三、审批与签署", 1)
sign = doc.add_table(rows=1, cols=3)
sign.style = "Table Grid"
for i, text in enumerate(("审批角色", "签署信息", "签署栏")):
    set_cell_shading(sign.rows[0].cells[i], LIGHT)
    p = sign.rows[0].cells[i].paragraphs[0]
    set_para(p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.add_run(text), bold=True, color=BLUE)
for role in [
    "资产所有/业务主管",
    "平台运营技术负责人",
    "网络安全负责人",
    "法务/合规负责人",
    "应急恢复负责人",
    "测试负责人/被授权人",
]:
    cells = sign.add_row().cells
    vals = (role, "姓名：________\n职务/工号：________", "签名/盖章：________\n日期时间：________")
    for i, val in enumerate(vals):
        p = cells[i].paragraphs[0]
        set_para(p, after=0)
        for j, line in enumerate(val.split("\n")):
            if j:
                p.add_run().add_break()
            set_font(p.add_run(line), bold=(i == 0))
set_table_geometry(sign, [2400, 3360, 3600])
mark_header_row(sign.rows[0])

add_heading(doc, "十四、附件清单", 1)
add_kv_table(doc, [
    ("附件 1", "指定测试资产清单（含实例 UUID/订单号/资源池/IP/磁盘范围）"),
    ("附件 2", "分阶段测试计划、命令清单、消息类别、速率与停止条件"),
    ("附件 3", "批准的合成 Windows OS/软件/KB/Agent 遥测数据集"),
    ("附件 4", "备份、快照、官方重装与人工恢复方案及演练记录"),
    ("附件 5", "后台观察指标、告警预期与验收记录表"),
    ("附件 6", "测试组件、镜像、脚本和最终交付物 SHA-256 清单"),
    ("附件 7", "保密、数据处理和漏洞披露要求"),
])

add_heading(doc, "签署前完整性检查", 1)
for text in [
    "□ 授权单位为完整法人名称，且签署人具有授权权限。",
    "□ 已填写唯一实例 UUID/订单号、资源池和允许改写的磁盘范围。",
    "□ 身份、遥测、heartbeat、协议构造/重放分别逐项勾选。",
    "□ 已批准合成数据集、消息类别、速率、凭据处理与销毁时间。",
    "□ 后台观察人和应急恢复负责人确认测试窗口并可实时联络。",
    "□ 已验证官方重装/快照/替代实例等恢复路径。",
    "□ 网络安全和法务/合规负责人已经签署。",
    "□ 所有附件与本授权书版本一致并已加盖骑缝章或电子签名。",
]:
    add_para(doc, text, after=3)

doc.core_properties.title = "云桌面 Debian 兼容性与 Guest 管理安全测试授权书范本"
doc.core_properties.subject = "受控操作系统替换、Guest 身份、管理协议与遥测仿真测试授权"
doc.core_properties.author = ""
doc.core_properties.keywords = "云桌面, Debian, Guest Agent, 安全测试, 授权书"
doc.save(OUT)
print(OUT)
